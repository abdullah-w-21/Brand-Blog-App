import streamlit as st
from docx import Document
from crewai import Agent, Task, Crew, Process
from langchain_openai import ChatOpenAI
from datetime import datetime
from io import BytesIO
import os

# Brand-specific information
BRAND_INFO = {
    "nao medical": {
        "description": "Nao Medical is a modern healthcare service provider focused on accessible, high-quality medical care",
        "voice": "Healthcare-focused, empathetic, and professional while maintaining accessibility",
        "keywords": ["healthcare", "medical care", "patient care", "accessible healthcare"]
    },
    "quixoa": {
        "description": "Quixoa is a powerful AI platform that brings together leading AI models like ChatGPT, Claude, and Gemini in one convenient place. Designed for businesses, Quixoa offers customizable, scalable, and affordable AI solutions tailored to your specific needs",
        "voice": "Tech-savvy, forward-thinking, and solution-oriented",
        "keywords": ["technology", "innovation", "digital solutions"]
    },
    "harmoni": {
        "description": "Harmoni is a cutting-edge medical translation service that provides seamless, HIPAA-compliant solutions for enhanced patient-provider communication. It's an AI-driven platform that offers real-time, accurate translations tailored to healthcare needs.",
        "voice": "Holistic, nurturing, and wellness-focused",
        "keywords": ["wellness", "lifestyle", "balance", "well-being"]
    },
    "youlogist": {
        "description": "Youologist is an innovative AI health companion designed to revolutionize your self-care journey. With 24/7 availability, personalized health insights, and real-time support, Youologist provides tailored guidance and clear health tips based on your unique health history.",
        "voice": "Empowering, motivational, and growth-oriented",
        "keywords": ["personal development", "growth", "coaching"]
    }
}

# Platform-specific settings
PLATFORM_SETTINGS = {
    "blog": {
        "min_words": 200,
        "max_words": 1500,
        "default_words": 400,
        "format_guidelines": """
        - Include SEO-optimized headings and subheadings
        - Use meta description-friendly introduction
        - Include internal and external linking opportunities
        - Break content into scannable sections
        """
    },
    "linkedin": {
        "min_words": 100,
        "max_words": 300,
        "default_words": 200,
        "format_guidelines": """
        - Start with a hook to capture attention
        - Use short, impactful paragraphs
        - Include relevant hashtags
        - End with a clear call-to-action
        - Format for mobile readability
        """
    }
}


class ResearchAgent(Agent):
    def __init__(self):
        super().__init__(
            role='Expert Content Researcher',
            goal='Research and analyze topic to gather key information tailored to brand voice and platform',
            backstory="""Expert researcher with deep knowledge of various industries, 
            brand voice adaptation, and platform-specific content strategy. Specializes in identifying 
            key trends and insights that align with specific brand identities and platform requirements.""",
            verbose=True,
            allow_delegation=False,
            llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
        )


class ContentWriter(Agent):
    def __init__(self):
        super().__init__(
            role='Brand-Focused Content Writer',
            goal='Write engaging, platform-optimized content in specified style and tone',
            backstory="""Expert content writer specializing in creating engaging, 
            well-researched content that perfectly matches brand voice, platform requirements,
            and style guidelines. Skilled in various writing styles and tones.""",
            verbose=True,
            allow_delegation=False,
            llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.7)
        )


def create_word_doc(content):
    doc = Document()
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('Title:'):
            doc.add_heading(line[6:].strip(), level=1)
        elif line.startswith('H2:'):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('H3:'):
            doc.add_heading(line[3:].strip(), level=3)
        elif line:
            doc.add_paragraph(line)
    return doc


def count_words(text):
    return len(text.split())


def generate_content(primary_keyword, brand_name, platform, narration_style, tone, additional_keywords, word_count,
                     additional_info=None):
    researcher = ResearchAgent()
    writer = ContentWriter()

    brand_context = BRAND_INFO[brand_name.lower()]
    platform_context = PLATFORM_SETTINGS[platform.lower()]

    style_guidelines = {
        "listicle": """Format as a numbered list with clear, distinct points. 
        Each point should have a clear subheading and detailed explanation.""",
        "paragraphs": """Format as flowing paragraphs with smooth transitions. 
        Focus on narrative flow and connecting ideas naturally."""
    }

    tone_guidelines = {
        "urgency": "Use compelling language that creates a sense of immediate need or action",
        "care": "Express warmth, empathy, and genuine concern for the reader's well-being",
        "professional": "Maintain a sophisticated, authoritative, yet accessible voice",
        "FOMO": "Create excitement and highlight what readers might miss out on",
        "emotional": "Connect with readers on an emotional level using evocative language",
        "conversational": "Write as if having a friendly chat with the reader",
        "informative": "Focus on providing clear, valuable information and insights",
        "formal": "Maintain a scholarly, polished tone while remaining clear",
        "humorous": "Incorporate appropriate wit and light-hearted elements",
        "friendly": "Create a warm, approachable atmosphere in the writing",
        "curious": "Foster a sense of discovery and exploration",
        "optimistic": "Maintain an upbeat, positive perspective throughout"
    }

    research_task = Task(
        description=f"""
        Research this topic thoroughly with specific focus on {brand_name}'s perspective 
        and {platform} platform requirements:

        Brand Context:
        - Description: {brand_context['description']}
        - Brand Voice: {brand_context['voice']}
        - Core Keywords: {', '.join(brand_context['keywords'])}

        Platform Requirements:
        - Platform: {platform}
        {platform_context['format_guidelines']}

        Topic Details:
        - Primary Keyword: {primary_keyword}
        - Additional Keywords: {', '.join(additional_keywords) if additional_keywords else 'None'}
        - Additional Context: {additional_info if additional_info else 'None provided'}

        Consider:
        1. How this topic aligns with {brand_name}'s values and {platform} audience
        2. Platform-specific content optimization
        3. Current trends relevant to topic, brand, and platform
        4. Competitor analysis in this space
        5. Unique value proposition for {brand_name}

        Format your response as:
        PLATFORM STRATEGY:
        - Content approach for {platform}
        - Audience engagement tactics

        BRAND ALIGNMENT:
        - How topic serves brand goals
        - Target audience intersection

        KEY POINTS:
        - Point 1 (brand-relevant)
        - Point 2 (brand-relevant)

        TRENDS:
        - Trend 1 with stats
        - Trend 2 with stats

        OUTLINE:
        1. Section 1
        2. Section 2
        """,
        expected_output="""Structured research document with platform-specific strategy, 
        brand-aligned insights, and content outline""",
        agent=researcher
    )

    writing_task = Task(
        description=f"""
        Write {platform}-optimized content based on this research:
        {{research_task.output}}

        Brand Guidelines:
        - Brand: {brand_name}
        - Brand Voice: {brand_context['voice']}
        - Brand Description: {brand_context['description']}

        Platform Requirements:
        - Platform: {platform}
        {platform_context['format_guidelines']}

        Style Requirements:
        - Narration Style: {narration_style}
        {style_guidelines[narration_style]}

        Tone Requirements:
        - Primary Tone: {tone}
        {tone_guidelines[tone]}

        Technical Requirements:
        1. Target word count: {word_count} words
        2. Primary keyword: {primary_keyword}
        3. Additional keywords: {', '.join(additional_keywords) if additional_keywords else 'None'}
        4. Additional context: {additional_info if additional_info else 'None'}

        Format:
        Title: [Engaging {platform}-optimized title with primary keyword]

        [Platform-appropriate introduction that establishes brand voice and tone]

        H2: [Section heading]
        [Section content matching selected style and tone]

        H2: [Section heading]
        [Section content matching selected style and tone]

        H3: [Subsection heading where relevant]
        [Subsection content]

        [Platform-optimized conclusion with brand-aligned call to action]
        """,
        expected_output=f"""Well-structured {platform} content that aligns with brand voice, 
        platform requirements, selected style, and tone while meeting all technical requirements""",
        agent=writer
    )

    crew = Crew(
        agents=[researcher, writer],
        tasks=[research_task, writing_task],
        verbose=True,
        process=Process.sequential
    )

    result = crew.kickoff()
    return str(result.tasks_output[1])


def main():
    st.title("Brand-Focused Content Generator")

    if 'blog_content' not in st.session_state:
        st.session_state.blog_content = None

    api_key = st.sidebar.text_input("OpenAI API Key", type="password")
    if not api_key:
        st.warning("Please enter your OpenAI API key")
        return

    os.environ["OPENAI_API_KEY"] = api_key

    with st.form("content_form"):
        # Platform and Brand Selection
        col1, col2 = st.columns(2)
        with col1:
            platform = st.selectbox(
                "Select Platform",
                options=list(PLATFORM_SETTINGS.keys()),
                format_func=lambda x: x.title()
            )
        with col2:
            brand_name = st.selectbox(
                "Select Brand",
                options=list(BRAND_INFO.keys()),
                format_func=lambda x: x.title()
            )

        # Style and Tone Selection
        col3, col4 = st.columns(2)
        with col3:
            narration_style = st.radio(
                "Narration Style",
                options=["listicle", "paragraphs"],
                format_func=lambda x: x.title()
            )
        with col4:
            tone = st.selectbox(
                "Content Tone",
                options=[
                    "urgency", "care", "professional", "FOMO", "emotional",
                    "conversational", "informative", "formal", "humorous",
                    "friendly", "curious", "optimistic"
                ],
                format_func=lambda x: x.title()
            )

        # Content Details
        primary_keyword = st.text_input("Primary Keyword (required)")
        additional_keywords = st.text_area("Additional Keywords (optional, one per line)")

        # Dynamic word count based on platform
        platform_settings = PLATFORM_SETTINGS[platform]
        word_count = st.number_input(
            "Word Count",
            min_value=platform_settings["min_words"],
            max_value=platform_settings["max_words"],
            value=platform_settings["default_words"],
            step=50
        )

        additional_info = st.text_area(
            "Additional Information (optional)",
            help="Add any specific points or context you want to include"
        )

        submitted = st.form_submit_button("Generate Content")

        if submitted and primary_keyword:
            with st.spinner(f"Generating {platform}-optimized content for {brand_name}..."):
                additional_keywords_list = [k.strip() for k in additional_keywords.split('\n') if k.strip()]

                try:
                    content = generate_content(
                        primary_keyword,
                        brand_name,
                        platform,
                        narration_style,
                        tone,
                        additional_keywords_list,
                        word_count,
                        additional_info
                    )
                    st.session_state.blog_content = content

                    actual_word_count = count_words(content)
                    st.success(f"Content generated successfully! Word count: {actual_word_count}")
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    print(f"Error details: {e}")

    if st.session_state.blog_content:
        st.subheader("Generated Content")
        st.markdown(st.session_state.blog_content)

        doc = create_word_doc(st.session_state.blog_content)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        st.download_button(
            "Download as Word Document",
            doc_buffer,
            f"content_{platform}_{timestamp}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


if __name__ == "__main__":
    main()